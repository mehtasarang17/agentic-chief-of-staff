"""Document processing for RAG system."""
import hashlib
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import uuid
from datetime import datetime

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings

from app.config import settings
from app.models.database import db_session, Document, DocumentChunk


class DocumentProcessor:
    """Process documents for RAG indexing."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.txt': 'text',
        '.md': 'markdown',
        '.docx': 'docx',
        '.doc': 'doc',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv'
    }

    def __init__(self):
        self.embeddings = OpenAIEmbeddings(
            model=settings.OPENAI_EMBEDDING_MODEL,
            api_key=settings.OPENAI_API_KEY
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    async def process_file(
        self,
        file_path: str,
        original_filename: str,
        user_id: Optional[str] = None,
        metadata: Dict[str, Any] = None
    ) -> Document:
        """
        Process a file and store it in the database with vector embeddings.

        Args:
            file_path: Path to the uploaded file
            original_filename: Original name of the file
            user_id: Optional user ID
            metadata: Additional metadata

        Returns:
            Document model instance
        """
        # Get file info
        file_ext = Path(original_filename).suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(file_ext, 'unknown')
        file_size = os.path.getsize(file_path)

        # Calculate file hash
        content_hash = self._calculate_hash(file_path)

        # Extract text content
        text_content = await self._extract_text(file_path, file_type)

        session = db_session()

        try:
            # Create document record
            document = Document(
                user_id=uuid.UUID(user_id) if user_id else None,
                filename=os.path.basename(file_path),
                original_filename=original_filename,
                file_type=file_type,
                file_size=file_size,
                file_path=file_path,
                content_hash=content_hash,
                processing_status='processing',
                metadata_=metadata or {}
            )
            session.add(document)
            session.flush()

            # Split into chunks
            chunks = self.text_splitter.split_text(text_content)

            # Generate embeddings for all chunks
            embeddings = self.embeddings.embed_documents(chunks)

            # Create chunk records
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx,
                    content=chunk_text,
                    embedding=embedding,
                    metadata_={
                        'chunk_size': len(chunk_text),
                        'original_filename': original_filename
                    }
                )
                session.add(chunk)

            # Update document status
            document.processing_status = 'completed'
            document.metadata_['chunk_count'] = len(chunks)
            document.metadata_['total_characters'] = len(text_content)

            session.commit()
            session.refresh(document)
            session.expunge(document)
            return document

        except Exception as e:
            session.rollback()
            # Update status to failed if document was created
            if 'document' in locals():
                document.processing_status = 'failed'
                document.metadata_['error'] = str(e)
                session.commit()
            raise e
        finally:
            session.close()

    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from file."""
        extractors = {
            'pdf': self._extract_pdf,
            'text': self._extract_text_file,
            'markdown': self._extract_text_file,
            'docx': self._extract_docx,
            'excel': self._extract_excel,
            'csv': self._extract_csv
        }

        extractor = extractors.get(file_type, self._extract_text_file)
        return await extractor(file_path)

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or '')
            text = '\n\n'.join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {e}")

        # Fallback for scanned or complex PDFs
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = [page.get_text("text") for page in doc]
            text = '\n\n'.join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {e}")

        return ""

    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {e}")

    async def _extract_excel(self, file_path: str) -> str:
        """Extract text from Excel file."""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, data_only=True)
            text_parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) for cell in row if cell])
                    if row_text:
                        text_parts.append(row_text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract Excel: {e}")

    async def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV file."""
        import csv
        text_parts = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(' | '.join(row))
        return '\n'.join(text_parts)

    def _calculate_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for byte_block in iter(lambda: f.read(4096), b''):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
